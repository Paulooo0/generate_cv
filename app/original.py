import subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()


def build_docx():
    # IMPORTANTE:
    # Aqui você deve manter o seu currículo original, que será utilizado como base para gerar o currículo final em build_docx.py.
    # O currículo original será mantido aqui para servir como base para que a IA gere o currículo final, sendo alimentada com o currículo original + a descrição da vaga no prompt.

    # Abaixo está um template de currículo, substitua as informações com as do seu currículo original.

    # Cabeçalho com nome e contatos
    add_centered_heading("Seu Nome Completo")
    add_centered_heading(
        "Cidade, Estado | (00) 00000-0000 | seu.email@exemplo.com",
        font_size=10,
    )
    add_centered_heading(
        "https://linkedin.com/in/seu-perfil | https://github.com/seuusuario",
        font_size=10,
    )

    # Seção: Resumo Profissional
    add_section_title("Resumo Profissional")
    add_paragraph_block(
        """Profissional da área de tecnologia com experiência em desenvolvimento de software, infraestrutura e automação. Atua com foco em boas práticas, qualidade de código e soluções escaláveis. Comunicativo, proativo e orientado a resultados."""
    )

    # Seção: Certificações (adicione ou remova conforme necessário)
    add_section_title("Certificações")
    add_paragraph_block(
        """Nome da Certificação – MM/AAAA
https://link-da-certificacao.com

Outra Certificação – MM/AAAA
https://link-da-certificacao.com"""
    )

    # Seção: Experiência Profissional
    add_section_title("Experiência Profissional")
    add_paragraph_block(
        """Nome da Empresa | Cargo | MM/AAAA – MM/AAAA
• Descreva suas responsabilidades e conquistas usando bullets.
• Foco em impacto, tecnologias usadas e resultados atingidos.

Outra Empresa | Cargo | MM/AAAA – Presente
• Repita o padrão para cada experiência relevante."""
    )

    # Seção: Educação
    add_section_title("Educação")
    add_paragraph_block(
        """Nome da Instituição – Curso | MM/AAAA – MM/AAAA

Outro Curso ou Treinamento | MM/AAAA – MM/AAAA"""
    )

    # Seção: Idiomas
    add_section_title("Idiomas")
    add_paragraph_block("Português – Nativo\nInglês – Avançado (C1)")

    # Seção: Habilidades Técnicas
    add_section_title("Habilidades Técnicas")
    add_paragraph_block(
        """Tecnologias: Python, Java, Node.js, Go, etc. 
Infraestrutura: AWS, Docker, Kubernetes, Terraform
Ferramentas: Git, Jenkins, CI/CD, Grafana
Banco de Dados: PostgreSQL, MongoDB, MySQL
Sistemas Operacionais: Linux, Windows
Outros: Metodologias ágeis (Scrum, Kanban), testes, segurança"""
    )

    return doc


def add_centered_heading(text, font_size=14):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)


def add_section_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.space_after = Pt(4)


def add_paragraph_block(text):
    for line in text.strip().split("\n"):
        doc.add_paragraph(line.strip())
